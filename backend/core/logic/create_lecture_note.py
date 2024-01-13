from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_lecture_note(file_path, dictionary):
    print("Creating lecture note...")

    doc = Document()

    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_table.autofit = False
    header_table.columns[0].width = Inches(2)

    # Title
    cell_title = header_table.cell(0, 0)
    title = cell_title.paragraphs[0].add_run(dictionary['Title'])
    title.bold = True
    title.font.size = Pt(17)
    cell_title.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('_' * 100, style='Normal')

    # Objectives
    objectives = doc.add_paragraph('\nObjectives of Lecture:', style='Normal')
    objectives.bold = True
    objectives.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = objectives.runs[0]
    run.font.size = Pt(17)

    for objective in dictionary["objective"]:
        doc.add_paragraph(objective, style='List Bullet')
    

    # Taught Things
    taught = doc.add_paragraph('\nTaught Things:', style='Normal')
    taught.bold = True
    taught.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = taught.runs[0]
    run.font.size = Pt(17)

    count = 1
    for context in dictionary["Taught Things"]:
        doc.add_paragraph(f'{count}. {context["Topic"]}', style='Normal')
        doc.add_paragraph(f'{context["Explanation of topic"]}', style='List Bullet')
        count += 1

    
    # Q and A
    q_a = doc.add_paragraph('\nQ & A:', style='Normal')
    q_a.bold = True
    q_a.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = q_a.runs[0]
    run.font.size = Pt(17)

    count = 1
    for qa in dictionary["Q/A from students"]:
        doc.add_paragraph(f'{count}. {qa["question"]}', style='Normal')
        doc.add_paragraph(f'{qa["answer"]}', style='List Bullet')
        count += 1

    # Conclusion
    conclusion = doc.add_paragraph('\nConclusion:', style='Normal')
    conclusion.bold = True
    conclusion.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = conclusion.runs[0]
    run.font.size = Pt(17)

    conclusion_t = doc.add_paragraph(f'\n{dictionary["Conclusion"]}', style='Normal')
    conclusion_t.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


    # Save the document
    doc.save(file_path)
    print(f'Minute created at: {file_path}')
    
    return None