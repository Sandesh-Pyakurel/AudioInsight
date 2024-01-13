from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_speech_document(file_path, dictionary, dictionary2):
    print("Creating speech document ...")
    
    doc = Document()

    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_table.autofit = False
    header_table.columns[0].width = Inches(2)

    # Title
    cell_title = header_table.cell(0, 0)
    title = cell_title.paragraphs[0].add_run(dictionary['Title'])
    title.bold = True
    title.font.size = Pt(17)
    cell_title.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    doc.add_paragraph('_' * 100, style='Normal')

    # Introduction
    introduction = doc.add_paragraph('\nIntroduction:', style='Normal')
    introduction.bold = True
    introduction.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = introduction.runs[0]
    run.font.size = Pt(17)

    introduction_text = doc.add_paragraph(f'\n{dictionary["Introduction"]}' + '\n', style='Normal')
    introduction_text.bold = True
    introduction_text.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


    # Speaker Experience
    experience = doc.add_paragraph('\nSpeaker Experience:', style='Normal')
    experience.bold = True
    experience.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = experience.runs[0]
    run.font.size = Pt(17)

    experience_text = doc.add_paragraph(f'\n{dictionary["Speaker Experience"]}' + '\n', style='Normal')
    experience_text.bold = True
    experience_text.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


    # Problems Shared
    problems = doc.add_paragraph('\nProblems Shared:', style='Normal')
    problems.bold = True
    problems.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = problems.runs[0]
    run.font.size = Pt(17)

    for problem in dictionary["Problems shared"]["description of problem"]:
        doc.add_paragraph(problem, style='List Bullet')

    
    # Solution
    solutions = doc.add_paragraph('\nSolutions:', style='Normal')
    solutions.bold = True
    solutions.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = solutions.runs[0]
    run.font.size = Pt(17)

    for solution in dictionary2["Solution"]:
        doc.add_paragraph(solution, style='List Bullet')

    #Current scenario
    if dictionary2["Current scenario"] != "":
        current_s = doc.add_paragraph('\nCurrent Scenario:', style='Normal')
        current_s.bold = True
        current_s.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = current_s.runs[0]
        run.font.size = Pt(17)
        
        current_s_text = doc.add_paragraph(f'\n{dictionary2["Current scenario"]}' + '\n', style='Normal')
        current_s_text.bold = True
        current_s_text.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Q and A
    qas = doc.add_paragraph('\nQ & A:', style='Normal')
    qas.bold = True
    qas.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = qas.runs[0]
    run.font.size = Pt(17)

    count = 1
    for qa in dictionary2["QandA from audience"]:
        doc.add_paragraph(f'{count}. {qa["question"]}', style='Normal')
        doc.add_paragraph(f'{qa["answer"]}', style='List Bullet')
        count += 1



    # Save the document
    doc.save(file_path)
    print(f'Minute created at: {file_path}')