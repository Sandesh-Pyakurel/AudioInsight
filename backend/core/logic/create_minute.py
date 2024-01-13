from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def suitable_text(value):
    if value == "":
        value = "HERE"
    return value


def create_minute(file_path, dictionary):
    print("Creating minute ...")

    doc = Document()

    header_table = doc.add_table(rows=4, cols=1)
    header_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_table.autofit = False
    header_table.columns[0].width = Inches(2)

    # Company Name
    cell_company_name = header_table.cell(0, 0)
    value = suitable_text(dictionary['Company name'])
    company_name = cell_company_name.paragraphs[0].add_run(value)
    company_name.bold = True
    company_name.font.size = Pt(17)
    company_name.font.color.rgb = RGBColor(255, 0, 0)
    cell_company_name.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Institute Name
    cell_institute = header_table.cell(1, 0)
    value = suitable_text(dictionary['Institute of'])
    institute_name = cell_institute.paragraphs[0].add_run(value)
    institute_name.bold = True
    institute_name.font.size = Pt(14)
    institute_name.font.color.rgb = RGBColor(255, 0, 0)
    cell_institute.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Address
    cell_address = header_table.cell(2, 0)
    value = suitable_text(dictionary['address'])
    address = cell_address.paragraphs[0].add_run(value)
    address.bold = True
    address.font.size = Pt(12)
    address.font.color.rgb = RGBColor(255, 0, 0)
    cell_address.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Contact Information
    cell_contact = header_table.cell(3, 0)
    value1, value2, value3 = suitable_text(dictionary["POB"]), suitable_text(dictionary["Ph. no."]), suitable_text(dictionary["Email"])
    contact_info = cell_contact.paragraphs[0].add_run( f'POB: {value1}\tPhone No.: {value2}\tEmail: {value3}' )
    contact_info.font.size = Pt(12)
    address.bold = True
    cell_contact.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    doc.add_paragraph('_' * 100, style='Normal')


    # Add date
    value = suitable_text(dictionary['Date of meeting'])
    date = doc.add_paragraph(f'Date: {value}\n', style='Normal')
    date.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


    # call to order
    call_to_order = doc.add_paragraph('Call to order', style='Normal')
    call_to_order.runs[0].bold = True
    call_to_order.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


    # start
    start = doc.add_paragraph(dictionary['starting text'] + '\n', style='Normal')
    start.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


    #Present members
    members_pre = doc.add_paragraph('Present Members:', style='Normal')
    members_pre.runs[0].bold = True
    members_pre.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    row_pre = 12
    col_pre = 4
    table_pre = doc.add_table(rows=row_pre, cols=col_pre)
    table_pre.style = 'Table Grid'
    
    widths_pre = (Inches(0.5), Inches(2.2), Inches(2.5), Inches(2.5))
    for row in table_pre.rows:
        for idx, width in enumerate(widths_pre):
            row.cells[idx].width = width

    column_names_pre = ["S.N.", "Present Member", "Position", "Signature"]
    for col_num, column_name in enumerate(column_names_pre):
        cell = table_pre.cell(0, col_num)
        cell.text = column_name
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(12)

    present_members = dictionary['present members']
    for row_num, row_name in enumerate(present_members):
        cell = table_pre.cell(row_num+1, 1)
        cell.text = row_name
        cell.paragraphs[0].runs[0].font.size = Pt(12)
        
        cell_sn = table_pre.cell(row_num+1, 0)
        cell_sn.text = f'{row_num+1}'
        cell_sn.paragraphs[0].runs[0].font.size = Pt(12)

    present_position = dictionary['present position']
    for row_num, row_name in enumerate(present_position):
        cell = table_pre.cell(row_num+1, 2)
        cell.text = row_name
        cell.paragraphs[0].runs[0].font.size = Pt(12)

    
    #Absent members
    members_abs = doc.add_paragraph('\nAbsent Members:', style='Normal')
    members_abs.runs[0].bold = True
    members_abs.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    row_abs = 5
    col_abs = 3
    table_abs = doc.add_table(rows=row_abs, cols=col_abs)
    table_abs.style = 'Table Grid'

    widths_abs = (Inches(0.5), Inches(2.2), Inches(2.5))
    for row in table_abs.rows:
        for idx, width in enumerate(widths_abs):
            row.cells[idx].width = width

    column_names_abs = ["S.N.", "Absent Member", "Position"]
    for col_num, column_name in enumerate(column_names_abs):
        cell = table_abs.cell(0, col_num)
        cell.text = column_name
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(12)

    absent_members = dictionary['absent members']
    for row_num, row_name in enumerate(absent_members):
        cell = table_abs.cell(row_num+1, 1)
        cell.text = row_name
        cell.paragraphs[0].runs[0].font.size = Pt(12)
        
        cell_sn = table_abs.cell(row_num+1, 0)
        cell_sn.text = f'{row_num+1}'
        cell_sn.paragraphs[0].runs[0].font.size = Pt(12)

    absent_position = dictionary['absent position']
    for row_num, row_name in enumerate(absent_position):
        cell = table_abs.cell(row_num+1, 2)
        cell.text = row_name
        cell.paragraphs[0].runs[0].font.size = Pt(12)
    

    #Agendas
    agendas = doc.add_paragraph('\nAGENDAS:' + '\n', style='Normal')
    agendas.bold = True
    agendas.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    agenda_count = 1
    for agenda_data in dictionary['Agendas']:
        title = agenda_data.get("title")
        description = agenda_data.get("description")

        agenda_tit = doc.add_paragraph(f'{agenda_count}. {title}:', style='Normal')
        agenda_tit.runs[0].bold = True
        agenda_tit.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        agenda_dis = doc.add_paragraph(f'{description}\n\n', style='Normal')
        agenda_dis.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        agenda_count += 1


    #ending
    conclusion = doc.add_paragraph(dictionary['conclude of meeting'] + '\n', style='Normal')
    conclusion.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    #signature
    signature = doc.add_paragraph('_'*14 + '\t\t\t\t\t\t\t\t' + '_'*14, style='Normal')
    signature.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    person = doc.add_paragraph("President" + '\t\t\t\t\t\t\t\t' + "Secretory", style='Normal')
    person.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Save the document
    doc.save(file_path)
    print(f'Minute created at: {file_path}')

